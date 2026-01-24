from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_line(doc, text, level=1):
    """Add a heading with a line underneath"""
    heading = doc.add_heading(text, level=level)
    heading.style = f'Heading {level}'

def add_table_of_contents(doc):
    """Add a table of contents"""
    doc.add_heading('Table of Contents', level=1)
    # Note: Word will auto-update this when document is opened
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('[Table of Contents will be generated when document is opened in Word]')
    run.italic = True
    doc.add_page_break()

def shade_paragraph(paragraph, color="D3D3D3"):
    """Add shading to a paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def add_code_block(doc, code, language=""):
    """Add a formatted code block"""
    paragraph = doc.add_paragraph()
    paragraph.style = 'Normal'
    shade_paragraph(paragraph, "F5F5F5")
    
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add left indent for code block
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)

# Create Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===== COVER PAGE =====
title = doc.add_heading('PETPAL MOBILE APPLICATION', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 102, 204)
title_run.font.size = Pt(28)
title_run.bold = True

doc.add_paragraph()

subtitle = doc.add_heading('Technical Report', level=1)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(20)

doc.add_paragraph()
doc.add_paragraph()

# Add metadata
info = doc.add_paragraph()
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info_run = info.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}\n')
info_run.font.size = Pt(12)
info_run = info.add_run('Version: 1.0\n')
info_run.font.size = Pt(12)
info_run = info.add_run('Repository: ammaribrahim95/MobileApp_PetPal')
info_run.font.size = Pt(12)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
add_table_of_contents(doc)

# ===== EXECUTIVE SUMMARY =====
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'PetPal is a comprehensive pet care management mobile application developed using Flutter and Firebase. '
    'The application connects pet owners with veterinarians, pet sitters, and hotel services while providing robust '
    'pet management, booking, and activity tracking capabilities. This report outlines the technical architecture, '
    'development approach, and implementation details of the PetPal application.'
)

# ===== DEVELOPMENT APPROACH =====
doc.add_heading('2. Development Approach and Workflow', level=1)

doc.add_heading('2.1 Architecture Pattern', level=2)
doc.add_paragraph(
    'The application follows a BLoC (Business Logic Component) architecture pattern, which provides:'
)
doc.add_paragraph('Clear separation of concerns between UI, business logic, and data layers', style='List Bullet')
doc.add_paragraph('Reactive state management using streams', style='List Bullet')
doc.add_paragraph('Testability and maintainability', style='List Bullet')
doc.add_paragraph('Predictable state transitions', style='List Bullet')

doc.add_heading('2.2 Project Structure', level=2)
doc.add_paragraph('The application follows a layered architecture:')

# Add structure visualization
structure = [
    ('PetPalApp Architecture', 'Root'),
    ('├── Presentation Layer (UI)', 'Layer'),
    ('│   ├── Screens (Views)', 'Component'),
    ('│   └── Widgets (Reusable Components)', 'Component'),
    ('├── Business Logic Layer', 'Layer'),
    ('│   ├── BLoCs (State Management)', 'Component'),
    ('│   └── Events & States', 'Component'),
    ('├── Data Layer', 'Layer'),
    ('│   ├── Repositories (Data Abstraction)', 'Component'),
    ('│   ├── Services (External APIs)', 'Component'),
    ('│   └── Models (Data Structures)', 'Component'),
    ('└── Core/Utils', 'Layer'),
    ('    ├── Constants', 'Component'),
    ('    └── Helpers', 'Component'),
]

for line, _ in structure:
    p = doc.add_paragraph(line, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_heading('2.3 Development Workflow', level=2)
workflow_items = [
    'Feature-Based Development: Each feature (authentication, pet management, bookings) is developed as an independent module',
    'Dependency Injection: All dependencies are injected through the widget tree using RepositoryProvider',
    'State Management: BLoC pattern ensures unidirectional data flow',
    'Firebase Integration: Backend-as-a-Service for authentication, database, storage, and notifications'
]
for i, item in enumerate(workflow_items, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_page_break()

# ===== SOFTWARE, TOOLS, AND FRAMEWORKS =====
doc.add_heading('3. Software, Tools, and Frameworks', level=1)

doc.add_heading('3.1 Core Technologies', level=2)

# Create table for technologies
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Technology'
hdr_cells[1].text = 'Version'
hdr_cells[2].text = 'Purpose'

tech_data = [
    ('Flutter', 'Latest Stable', 'Cross-platform mobile framework'),
    ('Dart', '3.0+', 'Programming language'),
    ('Firebase', 'Latest', 'Backend services'),
]

for i, (tech, version, purpose) in enumerate(tech_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = tech
    row_cells[1].text = version
    row_cells[2].text = purpose

doc.add_heading('3.2 Key Dependencies', level=2)

dependencies = """flutter_bloc: ^8.0.0           // BLoC pattern implementation
firebase_core: ^2.0.0           // Firebase initialization
firebase_auth: ^4.0.0           // Authentication
cloud_firestore: ^4.0.0         // NoSQL database
firebase_storage: ^11.0.0       // File storage
firebase_messaging: ^14.0.0     // Push notifications
flutter_local_notifications: ^15.0.0
pdf: PDF generation
image_picker: Image selection
intl: Internationalization and date formatting"""

add_code_block(doc, dependencies, "yaml")

doc.add_heading('3.3 Development Tools', level=2)
tools = [
    'IDE: Android Studio / Visual Studio Code',
    'Version Control: Git',
    'Firebase Console: Backend management',
    'Flutter DevTools: Debugging and profiling'
]
for tool in tools:
    doc.add_paragraph(tool, style='List Bullet')

doc.add_page_break()

# ===== DATABASE AND DATA HANDLING =====
doc.add_heading('4. Database and Data Handling Design', level=1)

doc.add_heading('4.1 Firebase Firestore Structure', level=2)
doc.add_paragraph('The application uses Cloud Firestore with the following collection structure:')

firestore_structure = """firestore
├── users/
│   ├── {userId}
│   │   ├── email: string
│   │   ├── name: string
│   │   ├── phoneNumber: string
│   │   ├── address: string
│   │   ├── profileImageUrl: string
│   │   └── role: string (owner/vet/sitter)
│
├── pets/
│   ├── {petId}
│   │   ├── ownerId: string
│   │   ├── name: string
│   │   ├── species: string
│   │   ├── breed: string
│   │   ├── age: number
│   │   ├── imageUrl: string
│   │   └── medicalHistory: array
│
├── bookings/
│   ├── {bookingId}
│   │   ├── ownerId: string
│   │   ├── petId: string
│   │   ├── providerId: string (vet/sitter)
│   │   ├── serviceType: string
│   │   ├── date: timestamp
│   │   ├── status: string
│   │   └── notes: string
│
├── vets/
│   ├── {vetId}
│   │   ├── userId: string
│   │   ├── clinicName: string
│   │   ├── specialization: string
│   │   ├── rating: number
│   │   └── availability: array
│
├── sitters/
│   ├── {sitterId}
│   │   ├── userId: string
│   │   ├── experience: string
│   │   ├── hourlyRate: number
│   │   └── availability: array
│
└── activities/
    ├── {activityId}
        ├── petId: string
        ├── type: string (feeding/walking/medication)
        ├── timestamp: timestamp
        └── notes: string"""

add_code_block(doc, firestore_structure)

doc.add_heading('4.2 Data Models', level=2)
doc.add_paragraph('Example AppUser Model:')

app_user_code = """class AppUser {
  final String id;
  final String email;
  final String name;
  final String? phoneNumber;
  final String? address;
  final String? profileImageUrl;
  final UserRole role;

  AppUser({
    required this.id,
    required this.email,
    required this.name,
    this.phoneNumber,
    this.address,
    this.profileImageUrl,
    required this.role,
  });

  Map<String, dynamic> toMap() {
    return {
      'id': id,
      'email': email,
      'name': name,
      'phoneNumber': phoneNumber,
      'address': address,
      'profileImageUrl': profileImageUrl,
      'role': role.toString(),
    };
  }

  factory AppUser.fromMap(Map<String, dynamic> map) {
    return AppUser(
      id: map['id'],
      email: map['email'],
      name: map['name'],
      phoneNumber: map['phoneNumber'],
      address: map['address'],
      profileImageUrl: map['profileImageUrl'],
      role: UserRole.values.firstWhere(
        (e) => e.toString() == map['role'],
      ),
    );
  }
}"""

add_code_block(doc, app_user_code, "dart")

doc.add_page_break()

# ===== CRUD OPERATIONS =====
doc.add_heading('5. Implementation of CRUD Operations', level=1)

doc.add_heading('5.1 Create Operation', level=2)
doc.add_paragraph('BLoC Event Handler for Adding a Pet:')

create_code = """on<AddPet>((event, emit) async {
  emit(state.copyWith(status: PetStatus.loading));
  
  try {
    String? imageUrl;
    
    if (event.imageFile != null) {
      imageUrl = await _storageService.uploadPetImage(
        event.pet.id,
        event.imageFile!,
      );
    }
    
    final pet = event.pet.copyWith(imageUrl: imageUrl);
    await _petRepository.addPet(pet);
    
    emit(state.copyWith(
      status: PetStatus.success,
      pets: [...state.pets, pet],
    ));
  } catch (e) {
    emit(state.copyWith(
      status: PetStatus.error,
      errorMessage: e.toString(),
    ));
  }
});"""

add_code_block(doc, create_code, "dart")

doc.add_heading('5.2 Read Operation', level=2)
doc.add_paragraph('Real-time Data Subscription:')

read_code = """on<LoadPets>((event, emit) async {
  emit(state.copyWith(status: PetStatus.loading));
  
  await emit.forEach<List<Pet>>(
    _petRepository.getPetsByOwner(event.ownerId),
    onData: (pets) => state.copyWith(
      status: PetStatus.success,
      pets: pets,
    ),
    onError: (error, stackTrace) => state.copyWith(
      status: PetStatus.error,
      errorMessage: error.toString(),
    ),
  );
});"""

add_code_block(doc, read_code, "dart")

doc.add_heading('5.3 Update Operation', level=2)
doc.add_paragraph('BLoC Handler for Updating a Pet:')

update_code = """on<UpdatePet>((event, emit) async {
  emit(state.copyWith(status: PetStatus.loading));
  
  try {
    String? imageUrl = event.pet.imageUrl;
    
    if (event.newImageFile != null) {
      if (imageUrl != null) {
        await _storageService.deleteFile(imageUrl);
      }
      
      imageUrl = await _storageService.uploadPetImage(
        event.pet.id,
        event.newImageFile!,
      );
    }
    
    final updatedPet = event.pet.copyWith(imageUrl: imageUrl);
    await _petRepository.updatePet(updatedPet);
    
    final updatedPets = state.pets.map((pet) => 
      pet.id == updatedPet.id ? updatedPet : pet
    ).toList();
    
    emit(state.copyWith(
      status: PetStatus.success,
      pets: updatedPets,
    ));
  } catch (e) {
    emit(state.copyWith(
      status: PetStatus.error,
      errorMessage: e.toString(),
    ));
  }
});"""

add_code_block(doc, update_code, "dart")

doc.add_heading('5.4 Delete Operation', level=2)
doc.add_paragraph('BLoC Handler for Deleting a Pet:')

delete_code = """on<DeletePet>((event, emit) async {
  emit(state.copyWith(status: PetStatus.loading));
  
  try {
    final pet = state.pets.firstWhere((p) => p.id == event.petId);
    
    if (pet.imageUrl != null) {
      await _storageService.deleteFile(pet.imageUrl!);
    }
    
    await _petRepository.deletePet(event.petId);
    
    final updatedPets = state.pets.where((p) => p.id != event.petId).toList();
    
    emit(state.copyWith(
      status: PetStatus.success,
      pets: updatedPets,
    ));
  } catch (e) {
    emit(state.copyWith(
      status: PetStatus.error,
      errorMessage: e.toString(),
    ));
  }
});"""

add_code_block(doc, delete_code, "dart")

doc.add_page_break()

# ===== IMPORTANT SOURCE CODE SNIPPETS =====
doc.add_heading('6. Important Source Code Snippets with Explanations', level=1)

doc.add_heading('6.1 Main Application Initialization', level=2)
doc.add_paragraph('From main.dart - Application entry point:')

main_code = """void main() async {
  WidgetsFlutterBinding.ensureInitialized();
  await Firebase.initializeApp(
    options: DefaultFirebaseOptions.currentPlatform
  );
  
  Bloc.observer = AppBlocObserver();

  final notificationService = NotificationService(
    FirebaseMessaging.instance,
    FlutterLocalNotificationsPlugin(),
  );

  final authService = FirebaseAuthService(FirebaseAuth.instance);
  final firestoreService = FirestoreService(FirebaseFirestore.instance);
  final storageService = StorageService(FirebaseStorage.instance);
  final pdfService = PdfService();

  final authRepository = AuthRepository(
    authService: authService,
    firestoreService: firestoreService,
  );
  
  // ... other repositories initialization

  runApp(PetPalApp(
    authRepository: authRepository,
    // ... other dependencies
  ));
}"""

add_code_block(doc, main_code, "dart")

doc.add_paragraph(
    'Explanation: Ensures Flutter bindings are initialized before Firebase, '
    'initializes all Firebase services and custom services, and creates repository instances '
    'with dependency injection for clean separation of concerns.'
)

doc.add_heading('6.2 Dependency Injection Setup', level=2)
doc.add_paragraph('Widget tree dependency configuration:')

di_code = """@override
Widget build(BuildContext context) {
  return MultiRepositoryProvider(
    providers: [
      RepositoryProvider.value(value: widget.authRepository),
      RepositoryProvider.value(value: widget.userRepository),
      RepositoryProvider.value(value: widget.petRepository),
      // ... other repositories
    ],
    child: MultiBlocProvider(
      providers: [
        BlocProvider(
          create: (_) => AuthBloc(widget.authRepository)
            ..add(const AuthStatusRequested()),
        ),
        BlocProvider(
          create: (_) => PetBloc(
            petRepository: widget.petRepository,
            storageService: widget.storageService,
          ),
        ),
        // ... other BLoCs
      ],
      child: MaterialApp(/* ... */),
    ),
  );
}"""

add_code_block(doc, di_code, "dart")

doc.add_paragraph(
    'Explanation: MultiRepositoryProvider makes repositories available throughout the widget tree. '
    'MultiBlocProvider creates and provides BLoC instances with their required dependencies. '
    'Initial events are dispatched on creation.'
)

doc.add_heading('6.3 Authentication Flow with BLoC Listener', level=2)

auth_listener_code = """BlocListener<AuthBloc, AuthState>(
  listenWhen: (previous, current) => previous.status != current.status,
  listener: (context, state) {
    if (state.status == AuthStatus.unauthenticated) {
      _navigatorKey.currentState?.pushNamedAndRemoveUntil(
        LoginScreen.routeName,
        (route) => false,
      );
    }
  },
  child: MaterialApp(
    navigatorKey: _navigatorKey,
    // ... routes
  ),
)"""

add_code_block(doc, auth_listener_code, "dart")

doc.add_paragraph(
    'Explanation: BlocListener monitors auth state changes without rebuilding UI. '
    'listenWhen prevents unnecessary listener executions. '
    'Global navigation key allows navigation from outside widget context, and '
    'pushNamedAndRemoveUntil clears the navigation stack when logging out.'
)

doc.add_page_break()

doc.add_heading('6.4 Firestore Service Implementation', level=2)

firestore_code = """class FirestoreService {
  final FirebaseFirestore _firestore;

  FirestoreService(this._firestore);

  CollectionReference collection(String path) {
    return _firestore.collection(path);
  }

  Future<DocumentSnapshot> getDocument(String path, String id) async {
    return await _firestore.collection(path).doc(id).get();
  }

  Stream<QuerySnapshot> getCollectionStream(
    String path,
    {List<WhereCondition>? where, String? orderBy}
  ) {
    Query query = _firestore.collection(path);

    if (where != null) {
      for (var condition in where) {
        query = query.where(condition.field, isEqualTo: condition.value);
      }
    }

    if (orderBy != null) {
      query = query.orderBy(orderBy);
    }

    return query.snapshots();
  }
}"""

add_code_block(doc, firestore_code, "dart")

doc.add_paragraph(
    'Explanation: Wraps FirebaseFirestore for easier testing and mocking. '
    'Provides generic CRUD operations, supports real-time streams for reactive UI, '
    'and allows query composition with where clauses and ordering.'
)

doc.add_heading('6.5 Storage Service for Image Handling', level=2)

storage_code = """class StorageService {
  final FirebaseStorage _storage;

  StorageService(this._storage);

  Future<String> uploadPetImage(String petId, File imageFile) async {
    try {
      final ref = _storage.ref()
          .child('pets/$petId/${DateTime.now().millisecondsSinceEpoch}');
      
      final uploadTask = ref.putFile(imageFile);
      final snapshot = await uploadTask.whenComplete(() {});
      
      final downloadUrl = await snapshot.ref.getDownloadURL();
      return downloadUrl;
    } catch (e) {
      throw Exception('Failed to upload image: $e');
    }
  }

  Future<void> deleteFile(String url) async {
    try {
      final ref = _storage.refFromURL(url);
      await ref.delete();
    } catch (e) {
      throw Exception('Failed to delete file: $e');
    }
  }
}"""

add_code_block(doc, storage_code, "dart")

doc.add_paragraph(
    'Explanation: Encapsulates Firebase Storage operations, generates unique file paths, '
    'returns download URLs for Firestore storage, and handles cleanup when images are updated/deleted.'
)

doc.add_page_break()

# ===== TECHNICAL CHALLENGES =====
doc.add_heading('7. Technical Challenges and Solutions', level=1)

challenges = [
    {
        'title': 'Asynchronous Initialization',
        'problem': 'Firebase and notification services require async initialization that could block app startup.',
        'solution': 'Initialize critical services (Firebase) synchronously in main(), then initialize non-critical services (notifications) asynchronously in widget\'s initState().',
        'outcome': 'Faster app startup with better user experience.'
    },
    {
        'title': 'State Management Complexity',
        'problem': 'Managing multiple interconnected states across different screens.',
        'solution': 'Implemented BLoC pattern with separate domain-specific BLoCs (AuthBloc, PetBloc, BookingBloc).',
        'outcome': 'Predictable state transitions, easy debugging, and testable business logic.'
    },
    {
        'title': 'Real-time Data Synchronization',
        'problem': 'Keeping UI synchronized with Firestore changes from multiple users/devices.',
        'solution': 'Use Firestore streams instead of futures. BLoCs subscribe to streams using emit.forEach().',
        'outcome': 'Automatic UI updates when Firestore data changes, no manual refresh required.'
    },
    {
        'title': 'Navigation After Logout',
        'problem': 'Users could navigate back to authenticated screens after logout.',
        'solution': 'Use pushNamedAndRemoveUntil() to clear entire navigation stack on logout.',
        'outcome': 'Complete navigation stack cleared, better security, consistent auth flow.'
    },
    {
        'title': 'Image Upload and Management',
        'problem': 'Handling image uploads, storage, and cleanup efficiently.',
        'solution': 'Encapsulate image operations in StorageService, delete old images before uploading new ones.',
        'outcome': 'Optimized storage usage, proper error handling, consistent image URLs.'
    },
    {
        'title': 'Role-Based Access Control',
        'problem': 'Different user types need access to different features.',
        'solution': 'Use UserRole enum, implement conditional routing and UI rendering based on role.',
        'outcome': 'Secure access control, customized user experience per role.'
    },
]

for i, challenge in enumerate(challenges, 1):
    doc.add_heading(f'7.{i} {challenge["title"]}', level=2)
    doc.add_paragraph(f'Problem: {challenge["problem"]}', style='List Bullet')
    doc.add_paragraph(f'Solution: {challenge["solution"]}', style='List Bullet')
    doc.add_paragraph(f'Outcome: {challenge["outcome"]}', style='List Bullet')

doc.add_page_break()

# ===== SECURITY CONSIDERATIONS =====
doc.add_heading('8. Security Considerations', level=1)

doc.add_heading('8.1 Authentication Security', level=2)
security_points = [
    'Firebase Authentication handles password hashing and secure token management',
    'Email verification for new accounts',
    'Password reset functionality with secure tokens'
]
for point in security_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('8.2 Data Access Rules', level=2)
doc.add_paragraph('Firestore security rules should enforce user-specific access:')

firestore_rules = """rules_version = '2';
service cloud.firestore {
  match /databases/{database}/documents {
    match /users/{userId} {
      allow read: if request.auth != null;
      allow write: if request.auth.uid == userId;
    }
    
    match /pets/{petId} {
      allow read: if request.auth != null;
      allow write: if request.auth != null && 
                    resource.data.ownerId == request.auth.uid;
    }
    
    match /bookings/{bookingId} {
      allow read: if request.auth != null && 
                 (resource.data.ownerId == request.auth.uid ||
                  resource.data.providerId == request.auth.uid);
    }
  }
}"""

add_code_block(doc, firestore_rules)

doc.add_heading('8.3 Storage Security', level=2)
doc.add_paragraph('Firebase Storage rules for authenticated access only:', style='List Bullet')
doc.add_paragraph('User-specific paths prevent unauthorized access', style='List Bullet')

doc.add_page_break()

# ===== PERFORMANCE OPTIMIZATIONS =====
doc.add_heading('9. Performance Optimizations', level=1)

doc.add_heading('9.1 Lazy Loading', level=2)
doc.add_paragraph('BLoCs created only when needed', style='List Bullet')
doc.add_paragraph('Images loaded on demand with caching', style='List Bullet')

doc.add_heading('9.2 Real-time Data Optimization', level=2)
doc.add_paragraph('Firestore queries with indexes for faster retrieval', style='List Bullet')
doc.add_paragraph('Limited collection queries with where clauses', style='List Bullet')
doc.add_paragraph('Pagination for large lists (recommended implementation)', style='List Bullet')

doc.add_heading('9.3 State Management Efficiency', level=2)
doc.add_paragraph('listenWhen and buildWhen prevent unnecessary rebuilds', style='List Bullet')
doc.add_paragraph('emit.forEach handles stream subscriptions efficiently', style='List Bullet')
doc.add_paragraph('Immutable state with copyWith for predictable updates', style='List Bullet')

doc.add_page_break()

# ===== TESTING STRATEGY =====
doc.add_heading('10. Testing Strategy', level=1)

doc.add_heading('10.1 Unit Testing', level=2)
doc.add_paragraph('BLoC business logic tested in isolation:', style='List Bullet')
doc.add_paragraph('Mock repositories and services', style='List Bullet')
doc.add_paragraph('Test state transitions and event handling', style='List Bullet')

unit_test_code = """void main() {
  group('PetBloc', () {
    late PetRepository mockRepository;
    late PetBloc petBloc;

    setUp(() {
      mockRepository = MockPetRepository();
      petBloc = PetBloc(petRepository: mockRepository);
    });

    test('emits success state when pets are loaded', () async {
      when(() => mockRepository.getPetsByOwner(any()))
          .thenAnswer((_) => Stream.value([mockPet]));

      expectLater(
        petBloc.stream,
        emitsInOrder([
          isA<PetState>().having((s) => s.status, 'status', PetStatus.loading),
          isA<PetState>().having((s) => s.status, 'status', PetStatus.success),
        ]),
      );

      petBloc.add(LoadPets(ownerId: 'test-owner'));
    });
  });
}"""

add_code_block(doc, unit_test_code, "dart")

doc.add_heading('10.2 Widget Testing', level=2)
doc.add_paragraph('UI components tested in isolation', style='List Bullet')
doc.add_paragraph('Mock BLoC states', style='List Bullet')
doc.add_paragraph('Test user interactions', style='List Bullet')

doc.add_heading('10.3 Integration Testing', level=2)
doc.add_paragraph('Complete user flows tested', style='List Bullet')
doc.add_paragraph('Firebase integration tested', style='List Bullet')
doc.add_paragraph('Navigation flows tested', style='List Bullet')

doc.add_page_break()

# ===== FUTURE ENHANCEMENTS =====
doc.add_heading('11. Future Enhancements', level=1)

doc.add_heading('11.1 Recommended Technical Improvements', level=2)

enhancements = [
    'Pagination: Implement pagination for large lists',
    'Offline Support: Enhanced offline capabilities with local database (Hive/Drift)',
    'Error Tracking: Integrate Crashlytics for production monitoring',
    'Analytics: Add Firebase Analytics for user behavior tracking',
    'Automated Testing: Increase test coverage to 80%+',
    'CI/CD: Set up automated build and deployment pipelines',
    'Performance Monitoring: Integrate Firebase Performance Monitoring',
    'Search Functionality: Implement Algolia or Elasticsearch',
    'Payment Integration: Add Stripe/PayPal for booking payments',
    'Chat Feature: Real-time messaging between owners and providers'
]

for i, enhancement in enumerate(enhancements, 1):
    doc.add_paragraph(f'{i}. {enhancement}')

doc.add_page_break()

# ===== CONCLUSION =====
doc.add_heading('12. Conclusion', level=1)

conclusion_text = (
    'PetPal demonstrates a well-architected Flutter application following industry best practices. '
    'The application successfully integrates multiple Firebase services (Authentication, Firestore, Storage, Messaging) '
    'with a robust state management solution, providing a solid foundation for a production-ready pet care management platform.\n\n'
    'Key Strengths:\n'
    '• Clean Architecture: Clear separation of concerns with BLoC pattern\n'
    '• Scalability: Modular design allows easy feature additions\n'
    '• Maintainability: Consistent code structure and patterns\n'
    '• Real-time Capabilities: Leveraging Firebase for real-time data sync\n'
    '• User Experience: Responsive UI with proper loading and error states\n'
    '• Security: Firebase Authentication and Firestore rules for data protection'
)

doc.add_paragraph(conclusion_text)

doc.add_paragraph()
doc.add_paragraph('---')
doc.add_paragraph('Report Prepared By: Technical Analysis Team')
doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
doc.add_paragraph('Version: 1.0')

# Save document
output_path = r'c:\Users\ammar\Documents\GitHub\MobileApp_PetPal\PetPal_Technical_Report.docx'
doc.save(output_path)
print(f'Document successfully created at: {output_path}')